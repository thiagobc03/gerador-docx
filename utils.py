from docx import Document
from docx.shared import Pt
import os

def gerar_docx(titulo, questoes, nome_arquivo):
    doc = Document()
    doc.add_heading(titulo, 0)

    for i, q in enumerate(questoes, start=1):
        doc.add_paragraph(f"{i}. {q['pergunta']}", style="List Number")
        for alt in q['alternativas']:
            doc.add_paragraph(f"   - {alt}")
        doc.add_paragraph(f"✅ Resposta correta: {q['respostaCorreta']}")
        doc.add_paragraph("")

    pasta = "./static"
    os.makedirs(pasta, exist_ok=True)
    caminho = os.path.join(pasta, nome_arquivo)
    doc.save(caminho)
    return caminho
